#!/usr/bin/env python3
"""LoveKorea.Us 개발 스펙 문서 생성 스크립트"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Noto Sans KR'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

NAVY = RGBColor(0x0B, 0x2E, 0x59)
RED = RGBColor(0xC8, 0x20, 0x2F)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = NAVY
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(18)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(14)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(12)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'💡 {text}')
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True
    return p


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LoveKorea.Us')
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('개발 스펙 문서 v1.0')
run.font.size = Pt(20)
run.font.color.rgb = RED

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Discover · Trust · Act')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.italic = True

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    '문서 유형: 개발 스펙 (Technical Specification)',
    '버전: 1.0',
    '작성일: 2026년 7월',
    '기반 문서: LoveKorea.Us 홈페이지 기획서 v1.1',
    '기술 스택: Next.js 15 + Supabase + Vercel',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('목차', level=1)
toc_items = [
    '1. 기술 아키텍처 (Tech Architecture)',
    '2. 데이터 모델 / DB 스키마',
    '3. API 엔드포인트 명세',
    '4. 인증/권한 설계 (Auth & RBAC)',
    '5. 페이지별 컴포넌트 트리',
    '6. MVP vs Phase 2+ 기능 분리',
    '7. 다국어(i18n) 설계',
    '8. SEO / 구조화 데이터 설계',
    '9. 프로젝트 구조 / 폴더 설계',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 1: 기술 아키텍처
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. 기술 아키텍처 (Tech Architecture)', level=1)

doc.add_heading('1.1 기술 스택 결정', level=2)
p = doc.add_paragraph('기획서 11장의 권장 스택을 기반으로, MVP 개발 속도와 운영 효율성을 고려하여 아래와 같이 최종 결정합니다.')

add_table(
    ['레이어', '기술', '선택 이유'],
    [
        ['프런트엔드', 'Next.js 15 (App Router)\nTypeScript 5.x\nTailwind CSS 4\nshadcn/ui', 'SSR/SSG/ISR 지원으로 SEO 최적화\nApp Router의 Server Components로 번들 최소화\nTailwind + shadcn으로 빠른 UI 개발\n기획서 권장 스택과 일치'],
        ['백엔드/DB', 'Supabase\n(PostgreSQL 15+)', 'PostgreSQL + Auth + Storage + Realtime 통합\nRow-Level Security(RLS)로 데이터 격리\nNext.js 공식 지원, Edge 호환\n무료 티어로 MVP 충분 (500MB DB, 1GB Storage)'],
        ['인증', 'Supabase Auth', 'DB와 자연 통합, RLS 연동\nEmail + Google + Kakao SSO 지원\nJWT 자동 관리, 미들웨어 통합\n별도 Auth 서비스 불필요'],
        ['검색', 'Algolia (MVP)\n→ Elasticsearch (Phase 3+)', '즉시 사용 가능한 SaaS, 한국어 형태소 지원\n자동완성, 필터링, 다국어 인덱스\nPhase 3에서 비용 최적화 위해 ES 마이그레이션'],
        ['영상', 'YouTube Data API v3\niframe (nocookie)', 'API로 자동 큐레이션 (키워드/채널 기반)\nlazy-load + nocookie로 성능/프라이버시\n일일 쿼터 10,000 유닛 (무료)'],
        ['파일 저장소', 'Supabase Storage\n+ Vercel Image Optimization', '상품/업체 이미지, 카탈로그 PDF\nVercel Edge에서 자동 WebP 변환, 리사이징'],
        ['배포/CDN', 'Vercel', 'Next.js 네이티브 지원, Edge Network\nPreview Deployments로 PR별 미리보기\nISR 자동 지원, 글로벌 CDN'],
        ['모니터링', 'Vercel Analytics\n+ Sentry', 'Core Web Vitals 실시간 추적\n에러 트래킹 및 알림'],
        ['이메일', 'Resend', '트랜잭션 이메일 (문의 알림, 인증)\nReact Email 템플릿 지원'],
    ],
    [5, 4.5, 8]
)

doc.add_heading('1.2 인프라 구성도', level=2)
p = doc.add_paragraph('아래는 전체 시스템 아키텍처의 개요입니다:')
doc.add_paragraph()

arch_text = """[Client (Browser/Mobile)]
        │
        ▼
[Vercel Edge Network] ─── CDN Cache (ISR)
        │
        ▼
[Next.js App Router]
   ├── Server Components (RSC) ── 직접 DB 쿼리
   ├── Route Handlers (/api/*) ── REST API
   └── Server Actions ── 폼 처리
        │
        ├──▶ [Supabase]
        │       ├── PostgreSQL (데이터)
        │       ├── Auth (인증/JWT)
        │       ├── Storage (파일)
        │       └── Realtime (알림)
        │
        ├──▶ [Algolia] ── 검색/자동완성
        │
        ├──▶ [YouTube API] ── 영상 큐레이션
        │
        └──▶ [Resend] ── 이메일 발송"""

add_code_block(arch_text)

doc.add_heading('1.3 렌더링 전략', level=2)
add_table(
    ['페이지', '렌더링', '캐시 전략', '근거'],
    [
        ['홈페이지', 'ISR', 'revalidate: 300 (5분)', '추천 상품/영상 주기적 갱신'],
        ['카테고리 리스트', 'SSG + ISR', 'revalidate: 600 (10분)', '카테고리 변경 빈도 낮음'],
        ['업체 상세', 'ISR', 'revalidate: 3600 (1시간)', '업체 정보 변경 빈도 낮음, SEO 중요'],
        ['상품 상세', 'ISR', 'revalidate: 1800 (30분)', '가격/재고 갱신 필요, SEO 중요'],
        ['Watch 피드', 'ISR', 'revalidate: 300', '영상 목록 주기적 갱신'],
        ['For Sellers', 'SSG', 'Build time', '정적 마케팅 페이지'],
        ['Deals', 'SSR', 'no-store', 'Flash Sale 실시간 카운트다운'],
        ['셀러 대시보드', 'CSR', 'SWR (client)', '인증 필요, 실시간 데이터'],
        ['검색 결과', 'CSR', 'Algolia cache', 'Algolia 클라이언트 검색'],
    ],
    [4, 3, 4.5, 6]
)

doc.add_heading('1.4 환경 변수', level=2)
add_table(
    ['변수명', '용도', '필수', '예시'],
    [
        ['NEXT_PUBLIC_SUPABASE_URL', 'Supabase 프로젝트 URL', 'Y', 'https://xxx.supabase.co'],
        ['NEXT_PUBLIC_SUPABASE_ANON_KEY', 'Supabase 공개 키', 'Y', 'eyJhbGci...'],
        ['SUPABASE_SERVICE_ROLE_KEY', 'Supabase 서비스 키 (서버용)', 'Y', 'eyJhbGci...'],
        ['NEXT_PUBLIC_ALGOLIA_APP_ID', 'Algolia 앱 ID', 'Y', 'XXXXXXXXXX'],
        ['NEXT_PUBLIC_ALGOLIA_SEARCH_KEY', 'Algolia 검색 키', 'Y', 'xxxx...'],
        ['ALGOLIA_ADMIN_KEY', 'Algolia 관리 키 (서버용)', 'Y', 'xxxx...'],
        ['YOUTUBE_API_KEY', 'YouTube Data API 키', 'Y', 'AIzaSy...'],
        ['RESEND_API_KEY', 'Resend 이메일 키', 'Y', 're_xxxx...'],
        ['NEXT_PUBLIC_SITE_URL', '사이트 URL', 'Y', 'https://lovekorea.us'],
    ],
    [5.5, 4, 1.5, 5]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 2: 데이터 모델
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. 데이터 모델 / DB 스키마', level=1)

doc.add_heading('2.1 ERD 개요', level=2)
p = doc.add_paragraph('Supabase(PostgreSQL) 기반 15개 핵심 테이블로 구성합니다. auth.users는 Supabase Auth가 관리하고, public 스키마에 비즈니스 테이블을 배치합니다.')

erd_text = """[auth.users] ──1:1──▶ [profiles]
                          │
                   ┌──────┴──────┐
                   ▼             ▼
          [seller_profiles]  [wishlists]
                   │             │
            ┌──────┼──────┐     │
            ▼      ▼      ▼     │
      [products] [videos] [certifications]
         │  │      │              │
         │  └──┬───┘              │
         │     ▼                  │
         │  [product_videos]      │
         │                        │
         ├──▶ [reviews] ◀────────┘
         │
         ├──▶ [leads]
         │
         ├──▶ [deals]
         │
         └──▶ [product_images]

[categories] ──self-ref──▶ parent_id
[subscriptions] ◀── seller_profiles
[ad_placements] ◀── seller_profiles
[translations] ◀── (polymorphic: products, sellers, categories)
[content_articles] ── 매거진/에디토리얼"""

add_code_block(erd_text)

doc.add_heading('2.2 테이블 상세', level=2)

# --- profiles ---
doc.add_heading('profiles (사용자 프로필)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK, FK→auth.users', '사용자 ID'],
        ['role', "enum('buyer','seller','admin')", "NOT NULL, DEFAULT 'buyer'", '역할'],
        ['display_name', 'text', 'NOT NULL', '표시 이름'],
        ['avatar_url', 'text', '', '프로필 이미지 URL'],
        ['preferred_locale', "varchar(5)", "DEFAULT 'en'", '선호 언어 (en/ko/ja/zh)'],
        ['country', 'varchar(2)', '', 'ISO 3166-1 alpha-2'],
        ['phone', 'varchar(20)', '', '연락처'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '가입일'],
        ['updated_at', 'timestamptz', 'DEFAULT now()', '수정일'],
    ],
    [3.5, 4, 4.5, 5]
)

# --- seller_profiles ---
doc.add_heading('seller_profiles (셀러/업체 프로필)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK, DEFAULT gen_random_uuid()', '셀러 ID'],
        ['user_id', 'uuid', 'FK→profiles, UNIQUE', '소유자'],
        ['company_name', 'text', 'NOT NULL', '업체명 (한국어)'],
        ['company_name_en', 'text', 'NOT NULL', '업체명 (영어)'],
        ['slug', 'varchar(100)', 'UNIQUE, NOT NULL', 'URL 슬러그'],
        ['logo_url', 'text', '', '로고 이미지'],
        ['cover_image_url', 'text', '', '커버 이미지'],
        ['description', 'text', '', '업체 소개 (한국어)'],
        ['description_en', 'text', '', '업체 소개 (영어)'],
        ['seller_type', "enum('manufacturer','brand','distributor','small_biz','service')", 'NOT NULL', '업체 유형'],
        ['business_type', "enum('b2b','b2c','both')", "DEFAULT 'both'", '거래 유형'],
        ['category_id', 'uuid', 'FK→categories', '주 카테고리'],
        ['target_markets', 'text[]', "DEFAULT '{}'", '진출 시장 (US/JP/CN/Global)'],
        ['website_url', 'text', '', '자사 웹사이트'],
        ['youtube_channel', 'text', '', 'YouTube 채널 URL'],
        ['contact_email', 'varchar(255)', '', '비즈니스 이메일'],
        ['contact_phone', 'varchar(20)', '', '비즈니스 전화'],
        ['address', 'jsonb', '', '주소 (country, city, detail)'],
        ['export_history', 'text', '', '수출 실적'],
        ['govt_support', 'boolean', 'DEFAULT false', '정부 지원사업 참여'],
        ['status', "enum('draft','pending','approved','rejected','suspended')", "DEFAULT 'draft'", '심사 상태'],
        ['is_verified', 'boolean', 'DEFAULT false', '인증 여부'],
        ['rating_avg', 'numeric(2,1)', 'DEFAULT 0', '평균 평점'],
        ['review_count', 'integer', 'DEFAULT 0', '리뷰 수'],
        ['subscription_tier', "enum('free','pro','premium')", "DEFAULT 'free'", '요금제'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '등록일'],
        ['updated_at', 'timestamptz', 'DEFAULT now()', '수정일'],
    ],
    [3.5, 5, 4, 5]
)

p = doc.add_paragraph('인덱스:')
for idx in [
    'idx_seller_status ON seller_profiles(status) WHERE status = \'approved\'',
    'idx_seller_category ON seller_profiles(category_id)',
    'idx_seller_markets ON seller_profiles USING GIN(target_markets)',
    'idx_seller_slug ON seller_profiles(slug)',
]:
    add_bullet(idx)

# --- categories ---
doc.add_heading('categories (카테고리)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '카테고리 ID'],
        ['parent_id', 'uuid', 'FK→categories, NULL', '상위 카테고리 (NULL=대분류)'],
        ['name', 'text', 'NOT NULL', '카테고리명 (영어)'],
        ['name_ko', 'text', '', '카테고리명 (한국어)'],
        ['slug', 'varchar(100)', 'UNIQUE', 'URL 슬러그'],
        ['icon', 'text', '', '아이콘 이름/URL'],
        ['sort_order', 'integer', 'DEFAULT 0', '정렬 순서'],
        ['depth', 'integer', 'DEFAULT 0', '깊이 (0=대, 1=중, 2=소)'],
    ],
    [3, 4, 5, 5]
)
p = doc.add_paragraph('기획서 5.1 기준 대분류 8개: Beauty & Cosmetics, Food & Health, Fashion & Lifestyle, Culture & Goods, Business & Trade, Directory, Watch, Deals & Ads')

# --- products ---
doc.add_heading('products (상품)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '상품 ID'],
        ['seller_id', 'uuid', 'FK→seller_profiles, NOT NULL', '판매 업체'],
        ['category_id', 'uuid', 'FK→categories', '카테고리'],
        ['name', 'text', 'NOT NULL', '상품명 (한국어)'],
        ['name_en', 'text', 'NOT NULL', '상품명 (영어)'],
        ['slug', 'varchar(150)', 'UNIQUE', 'URL 슬러그'],
        ['description', 'text', '', '상세 설명 (한국어)'],
        ['description_en', 'text', '', '상세 설명 (영어)'],
        ['price_min', 'numeric(12,2)', '', '최소 가격 (USD)'],
        ['price_max', 'numeric(12,2)', '', '최대 가격 (USD)'],
        ['moq', 'integer', '', '최소 주문 수량 (B2B)'],
        ['unit', 'varchar(20)', '', '단위 (pcs/box/kg 등)'],
        ['specs', 'jsonb', '', '스펙 (key-value)'],
        ['ingredients', 'text', '', '성분/원재료'],
        ['available_markets', 'text[]', "DEFAULT '{}'", '판매 가능 시장'],
        ['shipping_info', 'jsonb', '', '배송 정보'],
        ['purchase_url', 'text', '', '구매 링크 (제휴몰)'],
        ['is_sponsored', 'boolean', 'DEFAULT false', '스폰서 상품 여부'],
        ['status', "enum('draft','active','inactive')", "DEFAULT 'draft'", '상태'],
        ['rating_avg', 'numeric(2,1)', 'DEFAULT 0', '평균 평점'],
        ['review_count', 'integer', 'DEFAULT 0', '리뷰 수'],
        ['view_count', 'integer', 'DEFAULT 0', '조회수'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '등록일'],
        ['updated_at', 'timestamptz', 'DEFAULT now()', '수정일'],
    ],
    [3.5, 4, 4.5, 5]
)

# --- videos ---
doc.add_heading('videos (영상)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '영상 ID'],
        ['youtube_id', 'varchar(20)', 'UNIQUE, NOT NULL', 'YouTube 영상 ID'],
        ['title', 'text', 'NOT NULL', '영상 제목'],
        ['thumbnail_url', 'text', '', '썸네일 URL'],
        ['channel_name', 'text', '', '채널명'],
        ['channel_id', 'varchar(30)', '', 'YouTube 채널 ID'],
        ['duration', 'integer', '', '재생 시간 (초)'],
        ['view_count', 'integer', 'DEFAULT 0', '조회수'],
        ['video_type', "enum('review','factory','live','shorts','interview','other')", "DEFAULT 'review'", '영상 유형 (기획서 7.3 기준)'],
        ['source', "enum('api','seller','curated','owned')", "DEFAULT 'api'", '소싱 방식 (기획서 7.1 기준)'],
        ['seller_id', 'uuid', 'FK→seller_profiles, NULL', '연관 업체'],
        ['category_id', 'uuid', 'FK→categories, NULL', '카테고리'],
        ['is_featured', 'boolean', 'DEFAULT false', 'Watch 피처드'],
        ['published_at', 'timestamptz', '', '유튜브 게시일'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '등록일'],
    ],
    [3.5, 4.5, 4.5, 5]
)

# --- product_videos (junction) ---
doc.add_heading('product_videos (상품-영상 연결)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['product_id', 'uuid', 'FK→products', '상품'],
        ['video_id', 'uuid', 'FK→videos', '영상'],
        ['sort_order', 'integer', 'DEFAULT 0', '정렬'],
    ],
    [4, 4, 4, 5]
)
p = doc.add_paragraph('PK: (product_id, video_id)')

# --- product_images ---
doc.add_heading('product_images (상품 이미지)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '이미지 ID'],
        ['product_id', 'uuid', 'FK→products, NOT NULL', '상품'],
        ['url', 'text', 'NOT NULL', '이미지 URL (Supabase Storage)'],
        ['alt_text', 'text', '', '대체 텍스트'],
        ['sort_order', 'integer', 'DEFAULT 0', '정렬 순서'],
        ['is_primary', 'boolean', 'DEFAULT false', '대표 이미지'],
    ],
    [3, 4.5, 4.5, 5]
)

# --- reviews ---
doc.add_heading('reviews (리뷰)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '리뷰 ID'],
        ['user_id', 'uuid', 'FK→profiles, NOT NULL', '작성자'],
        ['product_id', 'uuid', 'FK→products, NULL', '상품 (NULL이면 업체 리뷰)'],
        ['seller_id', 'uuid', 'FK→seller_profiles, NULL', '업체'],
        ['rating', 'smallint', 'NOT NULL, CHECK(1-5)', '별점'],
        ['title', 'text', '', '리뷰 제목'],
        ['body', 'text', '', '리뷰 내용'],
        ['locale', 'varchar(5)', "DEFAULT 'en'", '작성 언어'],
        ['helpful_count', 'integer', 'DEFAULT 0', '유용해요 수'],
        ['status', "enum('active','flagged','removed')", "DEFAULT 'active'", '상태'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '작성일'],
    ],
    [3, 4.5, 4.5, 5]
)
p = doc.add_paragraph('제약: CHECK(product_id IS NOT NULL OR seller_id IS NOT NULL) — 최소 하나 필수')

# --- leads ---
doc.add_heading('leads (문의/견적)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '리드 ID'],
        ['buyer_id', 'uuid', 'FK→profiles, NULL', '바이어 (비회원이면 NULL)'],
        ['seller_id', 'uuid', 'FK→seller_profiles, NOT NULL', '수신 업체'],
        ['product_id', 'uuid', 'FK→products, NULL', '관련 상품'],
        ['type', "enum('inquiry','quote','partnership')", "DEFAULT 'inquiry'", '유형'],
        ['buyer_name', 'text', 'NOT NULL', '문의자 이름'],
        ['buyer_email', 'varchar(255)', 'NOT NULL', '문의자 이메일'],
        ['buyer_company', 'text', '', '문의자 회사'],
        ['buyer_country', 'varchar(2)', '', '국가'],
        ['message', 'text', 'NOT NULL', '문의 내용'],
        ['quantity', 'integer', '', '요청 수량'],
        ['status', "enum('new','read','replied','closed')", "DEFAULT 'new'", '상태'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '문의일'],
        ['replied_at', 'timestamptz', '', '응답일'],
    ],
    [3, 4, 4.5, 5]
)

# --- deals ---
doc.add_heading('deals (딜/프로모션)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '딜 ID'],
        ['product_id', 'uuid', 'FK→products, NOT NULL', '상품'],
        ['seller_id', 'uuid', 'FK→seller_profiles, NOT NULL', '업체'],
        ['deal_type', "enum('flash_sale','promo','sponsored','featured')", 'NOT NULL', '딜 유형'],
        ['title', 'text', 'NOT NULL', '딜 제목'],
        ['discount_percent', 'smallint', '', '할인율 (%)'],
        ['original_price', 'numeric(12,2)', '', '원래 가격'],
        ['deal_price', 'numeric(12,2)', '', '딜 가격'],
        ['starts_at', 'timestamptz', 'NOT NULL', '시작일'],
        ['ends_at', 'timestamptz', 'NOT NULL', '종료일'],
        ['is_active', 'boolean', 'DEFAULT true', '활성 여부'],
        ['sort_order', 'integer', 'DEFAULT 0', '노출 순서'],
    ],
    [3, 4, 4.5, 5]
)

# --- certifications ---
doc.add_heading('certifications (인증)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '인증 ID'],
        ['seller_id', 'uuid', 'FK→seller_profiles, NULL', '업체'],
        ['product_id', 'uuid', 'FK→products, NULL', '상품'],
        ['cert_type', "varchar(50)", 'NOT NULL', '인증 종류 (FDA/HACCP/ISO/GMP 등)'],
        ['cert_number', 'varchar(100)', '', '인증번호'],
        ['issued_by', 'text', '', '발급 기관'],
        ['valid_until', 'date', '', '유효기한'],
        ['document_url', 'text', '', '인증서 파일 URL'],
        ['is_verified', 'boolean', 'DEFAULT false', '관리자 검증 여부'],
    ],
    [3, 4, 4.5, 5]
)

# --- subscriptions ---
doc.add_heading('subscriptions (구독/요금제)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '구독 ID'],
        ['seller_id', 'uuid', 'FK→seller_profiles, UNIQUE', '업체'],
        ['tier', "enum('free','pro','premium')", "DEFAULT 'free'", '요금제'],
        ['status', "enum('active','past_due','cancelled')", "DEFAULT 'active'", '상태'],
        ['current_period_start', 'timestamptz', '', '현 결제 주기 시작'],
        ['current_period_end', 'timestamptz', '', '현 결제 주기 종료'],
        ['stripe_customer_id', 'varchar(50)', '', 'Stripe 고객 ID'],
        ['stripe_subscription_id', 'varchar(50)', '', 'Stripe 구독 ID'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '생성일'],
    ],
    [3.5, 4, 4.5, 5]
)

# --- ad_placements ---
doc.add_heading('ad_placements (광고 배치)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '광고 ID'],
        ['seller_id', 'uuid', 'FK→seller_profiles, NOT NULL', '광고주'],
        ['product_id', 'uuid', 'FK→products, NULL', '광고 상품'],
        ['placement', "enum('home_carousel','category_banner','search_top','newsletter')", 'NOT NULL', '노출 위치 (기획서 8.1 기준)'],
        ['billing_type', "enum('monthly','cpc','cpm')", 'NOT NULL', '과금 방식'],
        ['budget', 'numeric(10,2)', '', '예산'],
        ['spent', 'numeric(10,2)', 'DEFAULT 0', '소진액'],
        ['impressions', 'integer', 'DEFAULT 0', '노출수'],
        ['clicks', 'integer', 'DEFAULT 0', '클릭수'],
        ['starts_at', 'timestamptz', 'NOT NULL', '시작일'],
        ['ends_at', 'timestamptz', 'NOT NULL', '종료일'],
        ['is_active', 'boolean', 'DEFAULT true', '활성'],
    ],
    [3, 4.5, 4.5, 5]
)

# --- wishlists ---
doc.add_heading('wishlists (찜 목록)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['user_id', 'uuid', 'FK→profiles', '사용자'],
        ['product_id', 'uuid', 'FK→products', '상품'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '찜 날짜'],
    ],
    [4, 4, 4.5, 5]
)
p = doc.add_paragraph('PK: (user_id, product_id)')

# --- translations ---
doc.add_heading('translations (다국어 번역)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '번역 ID'],
        ['entity_type', "varchar(30)", 'NOT NULL', '대상 유형 (product/seller/category)'],
        ['entity_id', 'uuid', 'NOT NULL', '대상 ID'],
        ['locale', 'varchar(5)', 'NOT NULL', '언어 코드 (ja/zh)'],
        ['field', 'varchar(50)', 'NOT NULL', '필드명 (name/description)'],
        ['value', 'text', 'NOT NULL', '번역된 값'],
        ['is_auto', 'boolean', 'DEFAULT false', '자동 번역 여부'],
        ['updated_at', 'timestamptz', 'DEFAULT now()', '수정일'],
    ],
    [3, 4, 4, 5.5]
)
p = doc.add_paragraph('UNIQUE: (entity_type, entity_id, locale, field)')

# --- content_articles ---
doc.add_heading('content_articles (매거진/에디토리얼)', level=3)
add_table(
    ['컬럼', '타입', '제약', '설명'],
    [
        ['id', 'uuid', 'PK', '아티클 ID'],
        ['slug', 'varchar(200)', 'UNIQUE', 'URL 슬러그'],
        ['title', 'text', 'NOT NULL', '제목'],
        ['body', 'text', 'NOT NULL', '본문 (Markdown)'],
        ['cover_image_url', 'text', '', '커버 이미지'],
        ['category_id', 'uuid', 'FK→categories, NULL', '관련 카테고리'],
        ['author_id', 'uuid', 'FK→profiles', '작성자'],
        ['locale', 'varchar(5)', "DEFAULT 'en'", '작성 언어'],
        ['status', "enum('draft','published','archived')", "DEFAULT 'draft'", '상태'],
        ['published_at', 'timestamptz', '', '게시일'],
        ['created_at', 'timestamptz', 'DEFAULT now()', '생성일'],
    ],
    [3, 4.5, 4.5, 5]
)

doc.add_heading('2.3 RLS (Row-Level Security) 정책', level=2)
add_table(
    ['테이블', '정책', 'SELECT', 'INSERT', 'UPDATE', 'DELETE'],
    [
        ['profiles', '본인 프로필', '전체 공개', '자동 (Auth trigger)', '본인만', 'X'],
        ['seller_profiles', '셀러 프로필', 'approved만 공개', '인증 사용자', '소유자만', '소유자 (draft만)'],
        ['products', '상품', 'active만 공개', '소유 셀러', '소유 셀러', '소유 셀러'],
        ['videos', '영상', '전체 공개', '셀러 + Admin', '셀러 + Admin', 'Admin'],
        ['reviews', '리뷰', '전체 공개', '인증 사용자', '작성자', '작성자 + Admin'],
        ['leads', '문의', '당사자만', '전체(비회원 포함)', '수신 셀러', 'Admin'],
        ['deals', '딜', 'active만 공개', 'Admin', 'Admin', 'Admin'],
        ['wishlists', '찜', '본인만', '본인', '본인', '본인'],
        ['subscriptions', '구독', '본인 + Admin', 'System', 'System', 'X'],
        ['ad_placements', '광고', 'active 공개', '셀러 + Admin', 'Admin', 'Admin'],
    ],
    [2.5, 2, 2.5, 2.5, 2.5, 3]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 3: API 엔드포인트
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. API 엔드포인트 명세', level=1)

p = doc.add_paragraph('Next.js App Router Route Handlers와 Server Actions를 조합합니다. 공개 읽기 API는 Route Handler, 인증 필요 쓰기 작업은 Server Actions 우선 사용합니다.')

doc.add_heading('3.1 인증 (Auth)', level=2)
p = doc.add_paragraph('Supabase Auth SDK를 직접 사용하며, 별도 /api 라우트 없이 클라이언트에서 호출합니다.')
add_table(
    ['기능', '방식', '설명'],
    [
        ['회원가입', 'supabase.auth.signUp()', 'Email + Password'],
        ['로그인', 'supabase.auth.signInWithPassword()', 'Email 로그인'],
        ['Google SSO', 'supabase.auth.signInWithOAuth({provider:"google"})', 'Google 소셜 로그인'],
        ['Kakao SSO', 'supabase.auth.signInWithOAuth({provider:"kakao"})', '카카오 소셜 로그인'],
        ['로그아웃', 'supabase.auth.signOut()', '세션 종료'],
        ['비밀번호 재설정', 'supabase.auth.resetPasswordForEmail()', '이메일 발송'],
        ['세션 갱신', 'supabase.auth.getSession()', '자동 JWT 갱신'],
        ['Auth Callback', 'GET /api/auth/callback', 'OAuth 리다이렉트 처리'],
    ],
    [3.5, 6, 7]
)

doc.add_heading('3.2 셀러 API', level=2)
add_table(
    ['Method', 'Path / Action', '인증', 'Request', 'Response', '설명'],
    [
        ['GET', '/api/sellers', '-', 'category, market, type, page, limit', 'SellerList + Pagination', '셀러 목록 (필터)'],
        ['GET', '/api/sellers/[slug]', '-', '-', 'SellerDetail', '셀러 상세'],
        ['GET', '/api/sellers/[slug]/products', '-', 'page, limit', 'ProductList', '셀러의 상품 목록'],
        ['GET', '/api/sellers/[slug]/videos', '-', '-', 'VideoList', '셀러의 영상 목록'],
        ['POST', 'Server Action: createSeller', 'Buyer+', 'SellerFormData', 'SellerProfile', '셀러 등록 신청'],
        ['PUT', 'Server Action: updateSeller', 'Owner', 'SellerFormData', 'SellerProfile', '셀러 정보 수정'],
        ['GET', '/api/sellers/me/dashboard', 'Seller', '-', 'DashboardData', '셀러 대시보드 (리드/조회수/수익)'],
        ['GET', '/api/sellers/me/leads', 'Seller', 'status, page', 'LeadList', '수신 문의 목록'],
    ],
    [1.5, 5, 1.5, 3.5, 3, 3]
)

doc.add_heading('3.3 상품 API', level=2)
add_table(
    ['Method', 'Path / Action', '인증', 'Request', 'Response', '설명'],
    [
        ['GET', '/api/products', '-', 'category, market, minPrice, maxPrice, sort, page', 'ProductList + Pagination', '상품 목록 (필터/정렬)'],
        ['GET', '/api/products/[slug]', '-', '-', 'ProductDetail', '상품 상세 (영상/리뷰 포함)'],
        ['GET', '/api/products/featured', '-', 'limit', 'ProductList', '추천/스폰서 상품'],
        ['POST', 'Server Action: createProduct', 'Seller', 'ProductFormData', 'Product', '상품 등록'],
        ['PUT', 'Server Action: updateProduct', 'Owner', 'ProductFormData', 'Product', '상품 수정'],
        ['DELETE', 'Server Action: deleteProduct', 'Owner', 'productId', 'void', '상품 삭제 (soft)'],
    ],
    [1.5, 5, 1.5, 4.5, 3, 3]
)

doc.add_heading('3.4 영상 API', level=2)
add_table(
    ['Method', 'Path / Action', '인증', 'Request', 'Response', '설명'],
    [
        ['GET', '/api/videos', '-', 'category, type, featured, page', 'VideoList', 'Watch 피드'],
        ['GET', '/api/videos/[id]', '-', '-', 'VideoDetail + RelatedProducts', '영상 상세'],
        ['GET', '/api/videos/shorts', '-', 'limit', 'VideoList', 'K-Shorts (60초 이하)'],
        ['POST', 'Server Action: registerVideo', 'Seller', 'youtubeUrl, productIds', 'Video', '셀러 영상 등록'],
        ['POST', '/api/videos/curate (Admin)', 'Admin', 'youtubeIds, category', 'VideoList', '에디터 큐레이션'],
        ['POST', '/api/cron/youtube-sync', 'Cron', 'categories[]', 'SyncResult', 'YouTube API 자동 수집 (Phase 3)'],
    ],
    [1.5, 5, 1.5, 4.5, 3.5, 3]
)

doc.add_heading('3.5 문의/견적 API', level=2)
add_table(
    ['Method', 'Path / Action', '인증', 'Request', 'Response', '설명'],
    [
        ['POST', 'Server Action: submitLead', '-', 'LeadFormData', 'Lead', '문의/견적 제출 (비회원 가능)'],
        ['PUT', 'Server Action: replyLead', 'Seller', 'leadId, message', 'Lead', '문의 응답'],
        ['PUT', 'Server Action: updateLeadStatus', 'Seller', 'leadId, status', 'Lead', '상태 변경'],
    ],
    [1.5, 5, 1.5, 4.5, 2.5, 3.5]
)

doc.add_heading('3.6 리뷰 API', level=2)
add_table(
    ['Method', 'Path / Action', '인증', 'Request', 'Response', '설명'],
    [
        ['GET', '/api/reviews?productId=x', '-', 'productId/sellerId, sort, page', 'ReviewList', '리뷰 목록'],
        ['POST', 'Server Action: createReview', 'User', 'rating, title, body, productId/sellerId', 'Review', '리뷰 작성'],
        ['POST', 'Server Action: voteHelpful', 'User', 'reviewId', 'void', '유용해요'],
        ['DELETE', 'Server Action: deleteReview', 'Owner/Admin', 'reviewId', 'void', '리뷰 삭제'],
    ],
    [1.5, 5, 2, 5, 2.5, 2.5]
)

doc.add_heading('3.7 검색 API', level=2)
add_table(
    ['Method', 'Path / Action', '인증', 'Request', 'Response', '설명'],
    [
        ['GET', '/api/search', '-', 'q, category, market, type, page', 'SearchResults (products + sellers + videos)', '통합 검색'],
        ['GET', '/api/search/suggest', '-', 'q', 'Suggestions[]', '자동완성'],
    ],
    [1.5, 4, 1, 5, 4, 2.5]
)
p = doc.add_paragraph('검색은 Algolia 클라이언트 SDK에서 직접 호출하되, 서버사이드 인덱싱은 Supabase Webhook → Edge Function → Algolia Admin API 파이프라인으로 구성합니다.')

doc.add_heading('3.8 관리자 API', level=2)
add_table(
    ['Method', 'Path / Action', '인증', 'Request', 'Response', '설명'],
    [
        ['GET', '/api/admin/sellers/pending', 'Admin', 'page', 'SellerList', '심사 대기 셀러'],
        ['PUT', 'Server Action: approveSeller', 'Admin', 'sellerId, status, reason', 'Seller', '셀러 승인/거절'],
        ['GET', '/api/admin/stats', 'Admin', 'period', 'DashboardStats', '플랫폼 통계'],
        ['PUT', 'Server Action: manageDeal', 'Admin', 'DealFormData', 'Deal', '딜 관리'],
        ['PUT', 'Server Action: manageAd', 'Admin', 'AdFormData', 'AdPlacement', '광고 관리'],
    ],
    [1.5, 5, 1.5, 4, 3, 3]
)

doc.add_heading('3.9 공통 타입 정의 (TypeScript)', level=2)

types_code = """// === Core Types ===

type Locale = 'en' | 'ko' | 'ja' | 'zh'
type Market = 'US' | 'JP' | 'CN' | 'Global'
type SellerType = 'manufacturer' | 'brand' | 'distributor' | 'small_biz' | 'service'
type SellerStatus = 'draft' | 'pending' | 'approved' | 'rejected' | 'suspended'

interface Pagination {
  page: number
  limit: number
  total: number
  totalPages: number
}

interface ApiResponse<T> {
  data: T
  pagination?: Pagination
  error?: { code: string; message: string }
}

interface SellerSummary {
  id: string; slug: string; companyName: string; companyNameEn: string
  logoUrl: string; sellerType: SellerType; targetMarkets: Market[]
  categoryName: string; ratingAvg: number; reviewCount: number
  isVerified: boolean; badges: string[]
}

interface ProductSummary {
  id: string; slug: string; name: string; nameEn: string
  primaryImageUrl: string; priceMin: number; priceMax: number
  moq: number; sellerName: string; sellerSlug: string
  ratingAvg: number; reviewCount: number; isSponsored: boolean
}

interface VideoSummary {
  id: string; youtubeId: string; title: string; thumbnailUrl: string
  channelName: string; viewCount: number; videoType: string
  duration: number
}"""

add_code_block(types_code)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 4: Auth & RBAC
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. 인증/권한 설계 (Auth & RBAC)', level=1)

doc.add_heading('4.1 역할 정의', level=2)
add_table(
    ['역할', '코드', '획득 조건', '주요 권한'],
    [
        ['바이어', 'buyer', '회원가입 시 기본', '상품/업체 탐색, 리뷰 작성, 문의 보내기, 찜 목록'],
        ['셀러', 'seller', '셀러 등록 신청 → 관리자 승인', '바이어 권한 + 상품/영상 등록, 문의 수신/응답, 대시보드'],
        ['관리자', 'admin', '직접 DB 설정', '전체 권한 + 셀러 심사, 콘텐츠 관리, 광고/딜 관리, 통계'],
    ],
    [2, 2, 4.5, 8]
)

doc.add_heading('4.2 역할별 접근 매트릭스', level=2)
add_table(
    ['기능', '비회원', '바이어', '셀러', '관리자'],
    [
        ['상품/업체/영상 조회', 'O', 'O', 'O', 'O'],
        ['통합 검색', 'O', 'O', 'O', 'O'],
        ['문의/견적 보내기', 'O (이메일 필요)', 'O', 'O', 'O'],
        ['리뷰 작성', 'X', 'O', 'O', 'O'],
        ['찜 목록', 'X', 'O', 'O', 'O'],
        ['셀러 등록 신청', 'X', 'O', '-', 'O'],
        ['상품/영상 등록', 'X', 'X', 'O', 'O'],
        ['문의 수신/응답', 'X', 'X', 'O', 'O'],
        ['셀러 대시보드', 'X', 'X', 'O', 'O'],
        ['셀러 심사 (승인/거절)', 'X', 'X', 'X', 'O'],
        ['콘텐츠 큐레이션', 'X', 'X', 'X', 'O'],
        ['광고/딜 관리', 'X', 'X', 'X', 'O'],
        ['플랫폼 통계', 'X', 'X', 'X', 'O'],
    ],
    [4, 2.5, 2.5, 2.5, 2.5]
)

doc.add_heading('4.3 셀러 온보딩 상태 머신', level=2)
state_machine = """[회원가입] → buyer 역할 획득
      │
      ▼
[셀러 등록 신청] → status: 'draft' (입력 중)
      │
      ▼ (폼 제출)
[심사 대기] → status: 'pending'
      │           → 관리자에게 알림 (이메일 + 대시보드)
      │
      ├──▶ [승인] → status: 'approved', role: 'seller'
      │              → 셀러에게 승인 이메일 발송
      │              → 상품 등록/대시보드 접근 가능
      │
      └──▶ [거절] → status: 'rejected'
                    → 거절 사유 이메일 발송
                    → 재신청 가능 (수정 후)"""

add_code_block(state_machine)

doc.add_heading('4.4 미들웨어 설계', level=2)
middleware_code = """// middleware.ts
import { createServerClient } from '@supabase/ssr'
import { NextResponse } from 'next/server'

// 보호 경로 정의
const SELLER_ROUTES = ['/dashboard', '/products/new', '/products/edit']
const ADMIN_ROUTES  = ['/admin']

export async function middleware(request: NextRequest) {
  const supabase = createServerClient(...)
  const { data: { user } } = await supabase.auth.getUser()

  // 1) 인증 필요 경로 체크
  if (isProtectedRoute(request.url) && !user) {
    return NextResponse.redirect(new URL('/login', request.url))
  }

  // 2) 셀러 전용 경로
  if (isSellerRoute(request.url)) {
    const profile = await getProfile(user.id)
    if (profile.role !== 'seller' && profile.role !== 'admin') {
      return NextResponse.redirect(new URL('/sellers/register', request.url))
    }
  }

  // 3) 관리자 전용 경로
  if (isAdminRoute(request.url)) {
    const profile = await getProfile(user.id)
    if (profile.role !== 'admin') {
      return NextResponse.redirect(new URL('/', request.url))
    }
  }

  // 4) Locale 라우팅 (7장에서 상세)
  return intlMiddleware(request)
}"""

add_code_block(middleware_code)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 5: 컴포넌트 트리
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. 페이지별 컴포넌트 트리', level=1)

p = doc.add_paragraph('디자인 시안 8장을 기반으로 각 페이지의 컴포넌트를 분해합니다. 각 컴포넌트의 데이터 소스와 렌더링 전략을 명시합니다.')

# --- Homepage ---
doc.add_heading('5.1 Homepage ( / )', level=2)
add_table(
    ['컴포넌트', 'Props 주요 필드', '데이터 소스', '렌더링'],
    [
        ['GlobalHeader', 'locale, user, cartCount', 'Auth session', 'CSR'],
        ['LanguageToggle', 'currentLocale, locales[]', 'URL param', 'CSR'],
        ['SearchBar', 'placeholder, onSearch', '-', 'CSR (Algolia)'],
        ['HeroSection', 'headline, subtext, bgVideo, marketSelector', 'CMS / Static', 'SSG'],
        ['TrustBar', 'sellerCount, countryCount, videoCount', 'DB aggregate', 'ISR'],
        ['HowItWorks', 'steps: {icon, title, desc}[]', 'Static', 'SSG'],
        ['CategoryGrid', 'categories: {name, icon, slug, count}[]', 'DB categories', 'ISR'],
        ['FeaturedProducts', 'products: ProductSummary[], label', 'DB (is_sponsored / featured)', 'ISR'],
        ['WatchPreview', 'videos: VideoSummary[]', 'DB (is_featured)', 'ISR'],
        ['FeaturedBrands', 'sellers: SellerSummary[]', 'DB (is_verified, rating)', 'ISR'],
        ['SellerCTA', 'headline, benefits[], ctaUrl', 'Static', 'SSG'],
        ['GlobalFooter', 'categories, markets, socials, legal', 'Static + DB', 'SSG'],
    ],
    [3.5, 5, 4, 2.5]
)

# --- For Sellers ---
doc.add_heading('5.2 For Sellers ( /sellers )', level=2)
add_table(
    ['컴포넌트', 'Props 주요 필드', '데이터 소스', '렌더링'],
    [
        ['SellerHero', 'headline, stats, ctaUrl', 'Static + DB agg', 'ISR'],
        ['HowItWorksSteps', 'steps: {number, icon, title, desc}[]', 'Static', 'SSG'],
        ['PricingTable', 'plans: {name, price, features[], cta}[]', 'Static / Config', 'SSG'],
        ['VoucherSection', 'programs: {name, desc, amount, link}[]', 'CMS', 'ISR'],
        ['SuccessStories', 'stories: {seller, quote, metric, image}[]', 'DB content_articles', 'ISR'],
        ['SellerFAQ', 'items: {q, a}[]', 'Static', 'SSG'],
        ['FinalCTA', 'headline, subtext, ctaUrl', 'Static', 'SSG'],
    ],
    [3.5, 5, 4, 2.5]
)

# --- Watch ---
doc.add_heading('5.3 Watch ( /watch )', level=2)
add_table(
    ['컴포넌트', 'Props 주요 필드', '데이터 소스', '렌더링'],
    [
        ['WatchHero', 'featuredVideo: VideoSummary', 'DB (featured)', 'ISR'],
        ['CategoryTabs', 'categories[], activeCategory', 'DB categories', 'ISR'],
        ['MarketFilter', 'markets[], activeMarket', 'Static', 'CSR'],
        ['VideoGrid', 'videos: VideoSummary[], layout', 'DB videos', 'ISR'],
        ['VideoCard', 'video: VideoSummary', '-', '-'],
        ['KShortsRail', 'shorts: VideoSummary[]', 'DB (duration < 60)', 'ISR'],
        ['FeaturedChannel', 'channel: {name, avatar, subCount, videos[]}', 'YouTube API', 'ISR'],
        ['TrendingTags', 'tags: string[]', 'Algolia / DB', 'ISR'],
        ['LoadMore', 'onLoad, hasMore, isLoading', '-', 'CSR'],
    ],
    [3.5, 5, 4, 2.5]
)

# --- Directory ---
doc.add_heading('5.4 Directory ( /directory )', level=2)
add_table(
    ['컴포넌트', 'Props 주요 필드', '데이터 소스', '렌더링'],
    [
        ['DirectoryHero', 'totalSellers, searchBar', 'DB agg', 'ISR'],
        ['FilterSidebar', 'categories, markets, sellerTypes, badges', 'DB enums', 'CSR'],
        ['BusinessList', 'sellers: SellerSummary[], view: grid/list', 'DB + Algolia', 'CSR'],
        ['BusinessCard', 'seller: SellerSummary', '-', '-'],
        ['MapView (Phase 2)', 'sellers: {lat, lng, name}[]', 'DB addresses', 'CSR'],
        ['IndustryGrid', 'industries: {name, icon, count}[]', 'DB agg by category', 'ISR'],
        ['Pagination', 'page, totalPages, onPageChange', '-', 'CSR'],
    ],
    [3.5, 5, 4, 2.5]
)

# --- Product Detail ---
doc.add_heading('5.5 Product Detail ( /products/[slug] )', level=2)
add_table(
    ['컴포넌트', 'Props 주요 필드', '데이터 소스', '렌더링'],
    [
        ['Breadcrumb', 'path: {label, href}[]', 'Category hierarchy', 'ISR'],
        ['ProductGallery', 'images: {url, alt}[]', 'DB product_images', 'ISR'],
        ['ProductVideo', 'video: VideoSummary (inline player)', 'DB product_videos', 'CSR (lazy)'],
        ['ProductInfo', 'name, description, specs, ingredients', 'DB products', 'ISR'],
        ['PriceSection', 'priceMin, priceMax, moq, unit, businessType', 'DB products', 'ISR'],
        ['CertBadges', 'certs: {type, label, icon}[]', 'DB certifications', 'ISR'],
        ['ActionButtons', 'purchaseUrl, onInquiry, onWishlist', '-', 'CSR'],
        ['ProductTabs', 'specs, shipping, certDetails', 'DB products', 'ISR'],
        ['ReviewSection', 'reviews: Review[], avgRating, distribution', 'DB reviews', 'ISR + CSR'],
        ['ReviewForm', 'productId, onSubmit', '-', 'CSR'],
        ['RelatedProducts', 'products: ProductSummary[]', 'DB (same category)', 'ISR'],
        ['MoreVideos', 'videos: VideoSummary[]', 'DB product_videos', 'ISR'],
        ['SellerMiniCard', 'seller: SellerSummary', 'DB seller_profiles', 'ISR'],
    ],
    [3.5, 5.5, 4, 2]
)

# --- Brand Detail ---
doc.add_heading('5.6 Brand / Seller Detail ( /brands/[slug] )', level=2)
add_table(
    ['컴포넌트', 'Props 주요 필드', '데이터 소스', '렌더링'],
    [
        ['BrandHeader', 'logo, name, type, markets[], badges, rating', 'DB seller_profiles', 'ISR'],
        ['ActionSidebar', 'onInquiry, onQuote, catalogUrl, websiteUrl', '-', 'CSR'],
        ['VideoTabs', 'review/factory/live tabs, videos[]', 'DB videos', 'CSR (lazy)'],
        ['ProductCatalog', 'products: ProductSummary[], filters', 'DB products', 'ISR'],
        ['TrustCredentials', 'certs[], exportHistory, govtSupport, memberSince', 'DB certifications + seller', 'ISR'],
        ['BuyerReviews', 'reviews: Review[], stats', 'DB reviews', 'ISR'],
        ['ContactInfo', 'email, phone, address, map', 'DB seller_profiles', 'ISR'],
        ['SimilarBrands', 'sellers: SellerSummary[]', 'DB (same category)', 'ISR'],
    ],
    [3.5, 5.5, 4, 2]
)

# --- Deals ---
doc.add_heading('5.7 Deals & Ads ( /deals )', level=2)
add_table(
    ['컴포넌트', 'Props 주요 필드', '데이터 소스', '렌더링'],
    [
        ['DealsHero', 'headline, totalDeals', 'DB agg', 'SSR'],
        ['FlashSaleTimer', 'endsAt, title', 'DB deals (flash_sale)', 'CSR (timer)'],
        ['FlashSaleCards', 'deals: DealSummary[]', 'DB deals', 'SSR'],
        ['SponsoredGrid', 'products: ProductSummary[] (is_sponsored)', 'DB products', 'SSR'],
        ['FeaturedBrandsMonth', 'sellers: SellerSummary[]', 'DB ad_placements', 'SSR'],
        ['SellerPackages', 'packages: {name, price, features}[]', 'Static', 'SSG'],
        ['DealsByCategory', 'categories: {name, deals[]}[]', 'DB deals grouped', 'SSR'],
    ],
    [3.5, 5, 4, 2.5]
)

# --- Mobile ---
doc.add_heading('5.8 Mobile Home (반응형)', level=2)
add_table(
    ['컴포넌트', '데스크톱 대비 변경점', '비고'],
    [
        ['MobileHeader', '햄버거 메뉴 + 검색 아이콘', '로고 축소, GNB → 드로어'],
        ['MobileHero', '높이 축소, 텍스트 간소화', 'CTA 버튼 풀폭'],
        ['CategoryPills', '가로 스크롤 → 가로 스크롤 필', '터치 스와이프'],
        ['FeaturedScroll', '그리드 → 가로 캐러셀', '카드 크기 조정'],
        ['WatchPreview', '6개 → 4개 썸네일', '세로 2×2 그리드'],
        ['SellerBanner', '별도 배너 → 인라인 CTA', '간소화'],
        ['BottomNav', '데스크톱 없음 → 하단 네비게이션 바', 'Home/Search/Watch/My 4탭'],
    ],
    [3.5, 7, 5]
)
p = doc.add_paragraph('반응형 브레이크포인트: sm(640px), md(768px), lg(1024px), xl(1280px), 2xl(1536px). Tailwind CSS 기본값 사용.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 6: MVP vs Phase 2+
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. MVP vs Phase 2+ 기능 분리', level=1)

p = doc.add_paragraph('기획서 13장 로드맵을 기반으로, 시안에서 확인된 모든 기능을 Phase별로 분류합니다.')

doc.add_heading('6.1 Phase 1 — MVP (M2~M3)', level=2)
p = doc.add_paragraph('핵심 목표: "셀러가 등록하고, 바이어가 찾고 문의하는" 기본 루프 완성')

add_table(
    ['영역', '기능', '페이지/컴포넌트', '우선순위'],
    [
        ['홈', '히어로 + 카테고리 그리드 + 추천 상품 + Watch 프리뷰 + 셀러 CTA', 'Homepage', 'P0'],
        ['디렉토리', '업체 리스트 + 필터(카테고리/시장/유형) + 상세 페이지', 'Directory, Brand Detail', 'P0'],
        ['상품', '상품 리스트 + 필터 + 상세 페이지 (영상 임베드)', 'Product List, Product Detail', 'P0'],
        ['Watch', '영상 피드 + 카테고리 필터 + 상세 (연관 상품)', 'Watch, Video Detail', 'P0'],
        ['검색', 'Algolia 통합 검색 + 자동완성', 'SearchBar, SearchResults', 'P0'],
        ['셀러 등록', '가입 → 프로필 입력 → 심사 → 승인', 'Seller Registration Flow', 'P0'],
        ['상품 등록', '셀러가 상품 CRUD + 이미지 업로드', 'Product Form', 'P0'],
        ['문의/견적', '비회원도 가능한 문의 폼 + 셀러 알림', 'Lead Form', 'P0'],
        ['For Sellers', '요금제 안내 + HOW IT WORKS + CTA', 'For Sellers Landing', 'P0'],
        ['인증', 'Email + Google + Kakao SSO', 'Auth Pages', 'P0'],
        ['다국어', '영어 + 한국어 (UI 텍스트)', 'i18n (en, ko)', 'P0'],
        ['셀러 대시보드', '기본 통계(조회수/문의수)', 'Seller Dashboard', 'P1'],
        ['리뷰', '별점 + 텍스트 리뷰 (상품/업체)', 'Review Section', 'P1'],
        ['찜 목록', '바이어 관심 상품 저장', 'Wishlist', 'P1'],
        ['관리자', '셀러 심사 + 기본 통계', 'Admin Panel', 'P1'],
    ],
    [2, 5, 5, 2]
)

doc.add_heading('6.2 Phase 2 — 수익화 (M4~M5)', level=2)
add_table(
    ['영역', '기능', '비고'],
    [
        ['Deals 페이지', 'Flash Sale + 스폰서 상품 + 추천 브랜드', 'Deals Landing'],
        ['광고 상품', '홈 캐러셀 / 카테고리 배너 / 검색 상단 노출', '기획서 8.1 기준'],
        ['제휴 커미션', '구매 링크 → 제휴몰 연결 + 트래킹', 'Amazon/Coupang/Qoo10'],
        ['다국어 확장', '일본어 + 중국어 추가', 'i18n (ja, zh)'],
        ['뉴스레터', 'Resend 기반 주간 뉴스레터', 'Subscriber 수집'],
        ['셀러 요금제 결제', 'Stripe 연동 Pro/Premium 과금', 'subscriptions 테이블'],
        ['MapView', '디렉토리 지도 보기', 'Google Maps / Mapbox'],
        ['매거진/에디토리얼', '콘텐츠 아티클 발행', 'content_articles'],
    ],
    [3, 7, 5]
)

doc.add_heading('6.3 Phase 3 — 성장 (M6~M9)', level=2)
add_table(
    ['영역', '기능', '비고'],
    [
        ['YouTube API 자동화', 'Cron 기반 영상 자동 수집 + 태깅', '일일 쿼터 관리'],
        ['K-Shorts', '60초 이하 세로형 영상 피드', 'TikTok 스타일'],
        ['라이브 커머스', '라이브 스트리밍 임베드 + 실시간 채팅', 'YouTube Live API'],
        ['추천 알고리즘', '사용자 행동 기반 상품/영상 추천', 'Supabase + Edge Function'],
        ['Elasticsearch', 'Algolia → ES 마이그레이션 (비용 최적화)', '한국어 형태소 분석기'],
        ['A/B 테스트', 'Vercel Edge Config + Feature Flags', '전환율 최적화'],
        ['SNS 연동', 'Instagram/TikTok 피드 임베드', '옴니채널'],
    ],
    [3, 7, 5]
)

doc.add_heading('6.4 Phase 4 — 고도화 (M10~M12)', level=2)
add_table(
    ['영역', '기능', '비고'],
    [
        ['셀러 대시보드 고도화', '매출/트래픽 리포트, 광고 성과, ROI 분석', 'Recharts'],
        ['데이터 리포트', '월간 플랫폼 리포트 자동 생성', 'PDF 다운로드'],
        ['API 개방', '외부 파트너 연동 REST API', 'Rate limiting'],
        ['모바일 앱 (PWA)', 'Service Worker + Push 알림', 'next-pwa'],
        ['AI 챗봇', '바이어 상품 추천 챗봇', 'Vercel AI SDK'],
    ],
    [3, 7, 5]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 7: i18n
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. 다국어(i18n) 설계', level=1)

doc.add_heading('7.1 개요', level=2)
add_table(
    ['항목', '설계'],
    [
        ['라이브러리', 'next-intl (Next.js App Router 공식 권장)'],
        ['지원 언어', 'MVP: en(기본), ko | Phase 2: ja, zh'],
        ['URL 구조', '/en/products/..., /ko/products/... (prefix 방식)'],
        ['기본 언어', 'en (prefix 생략 가능: / = /en)'],
        ['감지 순서', 'URL prefix → Cookie → Accept-Language → 기본값(en)'],
        ['폴백', '번역 없으면 en 원문 표시 + 자동번역 라벨'],
    ],
    [3, 14]
)

doc.add_heading('7.2 콘텐츠 유형별 전략', level=2)
add_table(
    ['유형', '저장', '번역 방식', '예시'],
    [
        ['UI 텍스트\n(버튼, 라벨, 메뉴)', 'JSON 파일\n/messages/{locale}.json', '개발자가 직접 작성\nnext-intl useTranslations()', '"Add to Cart" → "장바구니"'],
        ['셀러/상품 콘텐츠\n(이름, 설명)', 'DB 기본 컬럼(ko/en)\n+ translations 테이블(ja/zh)', '셀러가 한/영 직접 입력\n일/중은 AI 자동번역 + 감수', '상품명, 업체 소개'],
        ['카테고리', 'categories 테이블\nname(en) + name_ko', '직접 관리\n(8개 대분류 + 소분류)', '"Beauty & Cosmetics" / "뷰티"'],
        ['매거진/에디토리얼', 'content_articles\nlocale 컬럼', '수동 번역\n(AI 초안 + 전문 감수)', '블로그 포스트'],
        ['이메일 템플릿', 'React Email\n+ locale 분기', '메시지 파일 참조', '문의 알림, 승인 안내'],
    ],
    [3, 3.5, 4, 5]
)

doc.add_heading('7.3 폴더 구조', level=2)
i18n_structure = """src/
├── i18n/
│   ├── config.ts          # locales, defaultLocale 설정
│   ├── request.ts         # getRequestConfig (next-intl)
│   └── navigation.ts      # Link, redirect, usePathname 래퍼
├── messages/
│   ├── en.json            # 영어 UI 텍스트
│   ├── ko.json            # 한국어 UI 텍스트
│   ├── ja.json            # 일본어 (Phase 2)
│   └── zh.json            # 중국어 (Phase 2)
└── app/
    └── [locale]/          # 다국어 라우팅 루트
        ├── layout.tsx     # locale 컨텍스트
        ├── page.tsx       # 홈
        ├── products/
        │   └── [slug]/
        └── ..."""

add_code_block(i18n_structure)

doc.add_heading('7.4 미들웨어 Locale 라우팅', level=2)
locale_middleware = """// middleware.ts (locale 부분)
import createMiddleware from 'next-intl/middleware'
import { locales, defaultLocale } from '@/i18n/config'

const intlMiddleware = createMiddleware({
  locales,               // ['en', 'ko', 'ja', 'zh']
  defaultLocale: 'en',
  localePrefix: 'as-needed',  // /en 생략, /ko/... 표시
  localeDetection: true,      // Accept-Language 자동 감지
})"""

add_code_block(locale_middleware)

doc.add_heading('7.5 DB 다국어 쿼리 패턴', level=2)
db_i18n = """// 상품 조회 시 다국어 처리 예시
async function getProduct(slug: string, locale: Locale) {
  const product = await supabase
    .from('products')
    .select('*, seller:seller_profiles(*), images:product_images(*)')
    .eq('slug', slug)
    .single()

  // locale에 따라 name/description 선택
  if (locale === 'en') {
    return { ...product, name: product.name_en, desc: product.description_en }
  }
  if (locale === 'ko') {
    return { ...product, name: product.name, desc: product.description }
  }

  // ja/zh: translations 테이블 조회
  const { data: trans } = await supabase
    .from('translations')
    .select('field, value')
    .match({ entity_type: 'product', entity_id: product.id, locale })

  const transMap = Object.fromEntries(trans.map(t => [t.field, t.value]))
  return {
    ...product,
    name: transMap.name || product.name_en,   // 폴백: 영어
    desc: transMap.description || product.description_en,
  }
}"""

add_code_block(db_i18n)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 8: SEO
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. SEO / 구조화 데이터 설계', level=1)

doc.add_heading('8.1 메타데이터 전략', level=2)
p = doc.add_paragraph('Next.js App Router의 generateMetadata 함수를 활용하여 페이지별 동적 메타 태그를 생성합니다.')

add_table(
    ['페이지', 'title 패턴', 'description', 'og:image'],
    [
        ['홈', 'LoveKorea.Us — Discover the Best of Korea', '플랫폼 소개 문구', '브랜드 OG 이미지'],
        ['카테고리', '{Category} — Korean Products | LoveKorea.Us', '카테고리 설명', '카테고리 대표 이미지'],
        ['상품 상세', '{상품명} by {업체명} | LoveKorea.Us', '상품 설명 첫 160자', '상품 대표 이미지'],
        ['업체 상세', '{업체명} — Korean {type} | LoveKorea.Us', '업체 소개 첫 160자', '업체 로고/커버'],
        ['Watch', 'Watch Korean Product Reviews | LoveKorea.Us', '영상 피드 설명', '대표 썸네일'],
        ['영상 상세', '{영상 제목} | LoveKorea.Us Watch', '영상 설명', '유튜브 썸네일'],
        ['Deals', 'Deals — Special Offers on Korean Products', '딜 설명', '딜 OG 이미지'],
    ],
    [3, 5, 4.5, 4]
)

doc.add_heading('8.2 구조화 데이터 (JSON-LD)', level=2)

doc.add_heading('Organization (홈)', level=3)
org_schema = """{
  "@context": "https://schema.org",
  "@type": "Organization",
  "name": "LoveKorea.Us",
  "url": "https://lovekorea.us",
  "logo": "https://lovekorea.us/logo.png",
  "sameAs": [
    "https://youtube.com/@lovekorea",
    "https://instagram.com/lovekorea.us"
  ]
}"""
add_code_block(org_schema)

doc.add_heading('WebSite + SearchAction (홈)', level=3)
site_schema = """{
  "@context": "https://schema.org",
  "@type": "WebSite",
  "name": "LoveKorea.Us",
  "url": "https://lovekorea.us",
  "potentialAction": {
    "@type": "SearchAction",
    "target": "https://lovekorea.us/search?q={search_term_string}",
    "query-input": "required name=search_term_string"
  }
}"""
add_code_block(site_schema)

doc.add_heading('Product (상품 상세)', level=3)
product_schema = """{
  "@context": "https://schema.org",
  "@type": "Product",
  "name": "{productName}",
  "image": ["{imageUrl1}", "{imageUrl2}"],
  "description": "{description}",
  "brand": { "@type": "Brand", "name": "{brandName}" },
  "offers": {
    "@type": "AggregateOffer",
    "lowPrice": "{priceMin}",
    "highPrice": "{priceMax}",
    "priceCurrency": "USD",
    "availability": "https://schema.org/InStock"
  },
  "aggregateRating": {
    "@type": "AggregateRating",
    "ratingValue": "{ratingAvg}",
    "reviewCount": "{reviewCount}"
  }
}"""
add_code_block(product_schema)

doc.add_heading('LocalBusiness (업체 상세)', level=3)
biz_schema = """{
  "@context": "https://schema.org",
  "@type": "LocalBusiness",
  "name": "{companyName}",
  "image": "{logoUrl}",
  "address": {
    "@type": "PostalAddress",
    "addressCountry": "KR",
    "addressLocality": "{city}"
  },
  "aggregateRating": {
    "@type": "AggregateRating",
    "ratingValue": "{ratingAvg}",
    "reviewCount": "{reviewCount}"
  }
}"""
add_code_block(biz_schema)

doc.add_heading('VideoObject (영상)', level=3)
video_schema = """{
  "@context": "https://schema.org",
  "@type": "VideoObject",
  "name": "{title}",
  "description": "{description}",
  "thumbnailUrl": "{thumbnailUrl}",
  "uploadDate": "{publishedAt}",
  "duration": "PT{minutes}M{seconds}S",
  "contentUrl": "https://youtube.com/watch?v={youtubeId}",
  "embedUrl": "https://www.youtube-nocookie.com/embed/{youtubeId}"
}"""
add_code_block(video_schema)

doc.add_heading('8.3 hreflang 태그', level=2)
hreflang_code = """// generateMetadata에서 자동 생성
export function generateMetadata({ params }) {
  return {
    alternates: {
      canonical: `https://lovekorea.us/${params.locale}/products/${params.slug}`,
      languages: {
        'en': `/en/products/${params.slug}`,
        'ko': `/ko/products/${params.slug}`,
        'ja': `/ja/products/${params.slug}`,   // Phase 2
        'zh': `/zh/products/${params.slug}`,   // Phase 2
        'x-default': `/en/products/${params.slug}`,
      }
    }
  }
}"""
add_code_block(hreflang_code)

doc.add_heading('8.4 사이트맵 / robots.txt', level=2)
p = doc.add_paragraph('Next.js App Router의 sitemap.ts와 robots.ts로 동적 생성합니다.')

sitemap_code = """// app/sitemap.ts — 동적 사이트맵
export default async function sitemap(): Promise<MetadataRoute.Sitemap> {
  const products = await getAllProductSlugs()
  const sellers  = await getAllSellerSlugs()
  const locales  = ['en', 'ko']

  const productUrls = products.flatMap(slug =>
    locales.map(locale => ({
      url: `https://lovekorea.us/${locale}/products/${slug}`,
      lastModified: new Date(),
      changeFrequency: 'weekly' as const,
      priority: 0.8,
    }))
  )

  const sellerUrls = sellers.flatMap(slug =>
    locales.map(locale => ({
      url: `https://lovekorea.us/${locale}/brands/${slug}`,
      lastModified: new Date(),
      changeFrequency: 'monthly' as const,
      priority: 0.7,
    }))
  )

  return [
    { url: 'https://lovekorea.us', priority: 1.0, changeFrequency: 'daily' },
    ...productUrls,
    ...sellerUrls,
  ]
}"""
add_code_block(sitemap_code)

doc.add_heading('8.5 성능 목표 (Core Web Vitals)', level=2)
add_table(
    ['지표', '목표', '전략'],
    [
        ['LCP (Largest Contentful Paint)', '< 2.5s', 'ISR + CDN, Image Optimization, Hero 이미지 priority'],
        ['CLS (Cumulative Layout Shift)', '< 0.1', '이미지/영상 aspect-ratio 고정, 폰트 size-adjust'],
        ['INP (Interaction to Next Paint)', '< 200ms', 'React Server Components, 최소 클라이언트 JS'],
        ['FCP (First Contentful Paint)', '< 1.8s', 'Edge Network + SSR, critical CSS inline'],
        ['TTFB (Time to First Byte)', '< 800ms', 'Vercel Edge, ISR cache hit'],
    ],
    [4, 2.5, 10]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CHAPTER 9: 프로젝트 구조
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. 프로젝트 구조 / 폴더 설계', level=1)

doc.add_heading('9.1 폴더 트리', level=2)

folder_tree = """lovekorea-us/
├── .env.local                    # 환경 변수 (gitignore)
├── .env.example                  # 환경 변수 예시
├── next.config.ts                # Next.js 설정
├── tailwind.config.ts            # Tailwind 설정
├── tsconfig.json                 # TypeScript 설정
├── package.json
│
├── public/
│   ├── images/                   # 정적 이미지 (로고, 아이콘)
│   ├── fonts/                    # 웹폰트 (Noto Sans KR)
│   └── og/                       # OG 이미지 템플릿
│
├── src/
│   ├── app/
│   │   ├── [locale]/             # ★ 다국어 라우팅 루트
│   │   │   ├── layout.tsx        # 루트 레이아웃 (Header/Footer)
│   │   │   ├── page.tsx          # 홈페이지
│   │   │   ├── products/
│   │   │   │   ├── page.tsx      # 상품 리스트
│   │   │   │   └── [slug]/
│   │   │   │       └── page.tsx  # 상품 상세
│   │   │   ├── brands/
│   │   │   │   └── [slug]/
│   │   │   │       └── page.tsx  # 업체 상세
│   │   │   ├── directory/
│   │   │   │   └── page.tsx      # 디렉토리
│   │   │   ├── watch/
│   │   │   │   ├── page.tsx      # Watch 피드
│   │   │   │   └── [id]/
│   │   │   │       └── page.tsx  # 영상 상세
│   │   │   ├── deals/
│   │   │   │   └── page.tsx      # Deals (Phase 2)
│   │   │   ├── sellers/
│   │   │   │   ├── page.tsx      # For Sellers 랜딩
│   │   │   │   └── register/
│   │   │   │       └── page.tsx  # 셀러 등록
│   │   │   ├── search/
│   │   │   │   └── page.tsx      # 검색 결과
│   │   │   ├── dashboard/        # ★ 셀러 대시보드 (인증 필요)
│   │   │   │   ├── layout.tsx
│   │   │   │   ├── page.tsx      # 대시보드 홈
│   │   │   │   ├── products/
│   │   │   │   │   ├── page.tsx  # 내 상품 관리
│   │   │   │   │   └── new/
│   │   │   │   │       └── page.tsx
│   │   │   │   ├── leads/
│   │   │   │   │   └── page.tsx  # 문의 관리
│   │   │   │   └── settings/
│   │   │   │       └── page.tsx  # 프로필 설정
│   │   │   ├── admin/            # ★ 관리자 패널 (Admin 전용)
│   │   │   │   ├── layout.tsx
│   │   │   │   ├── page.tsx
│   │   │   │   ├── sellers/
│   │   │   │   │   └── page.tsx  # 셀러 심사
│   │   │   │   └── content/
│   │   │   │       └── page.tsx  # 콘텐츠 관리
│   │   │   ├── login/
│   │   │   │   └── page.tsx
│   │   │   └── signup/
│   │   │       └── page.tsx
│   │   │
│   │   ├── api/
│   │   │   ├── auth/
│   │   │   │   └── callback/
│   │   │   │       └── route.ts  # OAuth 콜백
│   │   │   ├── sellers/
│   │   │   │   ├── route.ts      # GET: 셀러 목록
│   │   │   │   └── [slug]/
│   │   │   │       └── route.ts  # GET: 셀러 상세
│   │   │   ├── products/
│   │   │   │   ├── route.ts
│   │   │   │   └── [slug]/
│   │   │   │       └── route.ts
│   │   │   ├── videos/
│   │   │   │   └── route.ts
│   │   │   ├── reviews/
│   │   │   │   └── route.ts
│   │   │   ├── search/
│   │   │   │   ├── route.ts      # 통합 검색
│   │   │   │   └── suggest/
│   │   │   │       └── route.ts  # 자동완성
│   │   │   ├── admin/
│   │   │   │   └── ...
│   │   │   └── cron/
│   │   │       └── youtube-sync/
│   │   │           └── route.ts  # Phase 3
│   │   │
│   │   ├── sitemap.ts            # 동적 사이트맵
│   │   └── robots.ts             # robots.txt
│   │
│   ├── components/
│   │   ├── ui/                   # ★ shadcn/ui 컴포넌트
│   │   │   ├── button.tsx
│   │   │   ├── card.tsx
│   │   │   ├── dialog.tsx
│   │   │   ├── input.tsx
│   │   │   ├── select.tsx
│   │   │   ├── tabs.tsx
│   │   │   └── ...
│   │   ├── layout/               # ★ 레이아웃 컴포넌트
│   │   │   ├── global-header.tsx
│   │   │   ├── global-footer.tsx
│   │   │   ├── mobile-nav.tsx
│   │   │   ├── language-toggle.tsx
│   │   │   └── breadcrumb.tsx
│   │   ├── home/                 # ★ 홈 전용 컴포넌트
│   │   │   ├── hero-section.tsx
│   │   │   ├── trust-bar.tsx
│   │   │   ├── category-grid.tsx
│   │   │   ├── featured-products.tsx
│   │   │   ├── watch-preview.tsx
│   │   │   └── seller-cta.tsx
│   │   ├── product/              # ★ 상품 관련
│   │   │   ├── product-card.tsx
│   │   │   ├── product-gallery.tsx
│   │   │   ├── product-info.tsx
│   │   │   ├── price-section.tsx
│   │   │   └── product-form.tsx
│   │   ├── seller/               # ★ 셀러/업체 관련
│   │   │   ├── seller-card.tsx
│   │   │   ├── brand-header.tsx
│   │   │   ├── trust-credentials.tsx
│   │   │   └── seller-form.tsx
│   │   ├── video/                # ★ 영상 관련
│   │   │   ├── video-card.tsx
│   │   │   ├── video-player.tsx
│   │   │   ├── video-grid.tsx
│   │   │   └── shorts-rail.tsx
│   │   ├── review/               # ★ 리뷰
│   │   │   ├── review-section.tsx
│   │   │   ├── review-card.tsx
│   │   │   └── review-form.tsx
│   │   ├── search/               # ★ 검색
│   │   │   ├── search-bar.tsx
│   │   │   └── search-results.tsx
│   │   ├── lead/                 # ★ 문의
│   │   │   └── lead-form.tsx
│   │   └── shared/               # ★ 공용 컴포넌트
│   │       ├── cert-badges.tsx
│   │       ├── market-badge.tsx
│   │       ├── rating-stars.tsx
│   │       ├── pagination.tsx
│   │       └── how-it-works.tsx
│   │
│   ├── lib/
│   │   ├── supabase/
│   │   │   ├── client.ts         # 브라우저용 클라이언트
│   │   │   ├── server.ts         # 서버 컴포넌트용
│   │   │   ├── middleware.ts     # 미들웨어용
│   │   │   └── admin.ts         # Service Role 클라이언트
│   │   ├── algolia/
│   │   │   ├── client.ts         # 검색 클라이언트
│   │   │   └── indexer.ts        # 인덱싱 유틸
│   │   ├── youtube/
│   │   │   └── client.ts         # YouTube API 래퍼
│   │   ├── email/
│   │   │   └── send.ts           # Resend 래퍼
│   │   └── utils.ts              # 공용 유틸리티
│   │
│   ├── actions/                  # ★ Server Actions
│   │   ├── auth.ts
│   │   ├── seller.ts
│   │   ├── product.ts
│   │   ├── video.ts
│   │   ├── lead.ts
│   │   └── review.ts
│   │
│   ├── hooks/                    # 클라이언트 훅
│   │   ├── use-auth.ts
│   │   ├── use-wishlist.ts
│   │   └── use-search.ts
│   │
│   ├── types/                    # TypeScript 타입
│   │   ├── database.ts           # Supabase 자동생성 타입
│   │   ├── api.ts                # API Request/Response
│   │   └── index.ts              # 공용 타입
│   │
│   ├── i18n/                     # 다국어 설정
│   │   ├── config.ts
│   │   ├── request.ts
│   │   └── navigation.ts
│   │
│   └── messages/                 # UI 번역 파일
│       ├── en.json
│       └── ko.json
│
├── supabase/
│   ├── migrations/               # DB 마이그레이션
│   │   ├── 001_initial_schema.sql
│   │   ├── 002_rls_policies.sql
│   │   └── 003_seed_categories.sql
│   └── config.toml               # Supabase 로컬 설정
│
└── tests/                        # 테스트 (Vitest)
    ├── unit/
    └── e2e/                      # Playwright"""

add_code_block(folder_tree)

doc.add_heading('9.2 핵심 설정 파일', level=2)

doc.add_heading('next.config.ts', level=3)
next_config = """import createNextIntlPlugin from 'next-intl/plugin'

const withNextIntl = createNextIntlPlugin()

const nextConfig = {
  images: {
    remotePatterns: [
      { hostname: '*.supabase.co' },
      { hostname: 'img.youtube.com' },
      { hostname: 'i.ytimg.com' },
    ],
  },
  experimental: {
    serverActions: { bodySizeLimit: '2mb' },
  },
}

export default withNextIntl(nextConfig)"""
add_code_block(next_config)

doc.add_heading('Supabase 클라이언트 (서버)', level=3)
supabase_server = """// lib/supabase/server.ts
import { createServerClient } from '@supabase/ssr'
import { cookies } from 'next/headers'

export async function createClient() {
  const cookieStore = await cookies()
  return createServerClient(
    process.env.NEXT_PUBLIC_SUPABASE_URL!,
    process.env.NEXT_PUBLIC_SUPABASE_ANON_KEY!,
    {
      cookies: {
        getAll() { return cookieStore.getAll() },
        setAll(cookiesToSet) {
          cookiesToSet.forEach(({ name, value, options }) =>
            cookieStore.set(name, value, options)
          )
        },
      },
    }
  )
}"""
add_code_block(supabase_server)

doc.add_heading('9.3 개발 환경 셋업', level=2)
setup_steps = [
    ('1. 저장소 클론', 'git clone <repo> && cd lovekorea-us'),
    ('2. 패키지 설치', 'pnpm install'),
    ('3. 환경 변수', 'cp .env.example .env.local → 값 채우기'),
    ('4. Supabase 로컬', 'npx supabase start → DB + Auth 로컬 실행'),
    ('5. DB 마이그레이션', 'npx supabase db push'),
    ('6. 시드 데이터', 'npx supabase db seed'),
    ('7. 개발 서버', 'pnpm dev → http://localhost:3000'),
]
add_table(
    ['단계', '명령어'],
    [[s, c] for s, c in setup_steps],
    [4, 12]
)

doc.add_heading('9.4 개발 컨벤션', level=2)
conventions = [
    '패키지 매니저: pnpm (lockfile 일관성)',
    '코드 포맷: Biome (ESLint + Prettier 대체, 속도 우선)',
    '커밋 메시지: Conventional Commits (feat: / fix: / chore:)',
    '브랜치 전략: main (프로덕션) ← develop ← feature/*',
    '컴포넌트 네이밍: PascalCase 파일명, kebab-case 폴더명',
    'Server Component 기본, "use client" 최소화',
    'Supabase 타입: npx supabase gen types typescript > src/types/database.ts',
]
for c in conventions:
    add_bullet(c)

doc.add_paragraph()
doc.add_paragraph()

# ── Final note ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 본 개발 스펙 문서 끝 · LoveKorea.Us —')
run.font.color.rgb = GRAY
run.italic = True


# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
output_path = os.path.expanduser('~/Downloads/LoveKoreaUs_개발스펙_v1.0.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')
