#!/usr/bin/env python3
"""Insert Chapter 12 into LoveKorea.Us 기획서 docx."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re

INPUT  = "/Users/martin/Downloads/LoveKoreaUs_홈페이지_기획서.docx"
OUTPUT = "/Users/martin/Downloads/LoveKoreaUs_홈페이지_기획서_v1.1.docx"

doc = Document(INPUT)

# --- Find insertion point: after Ch 11 content, before Ch 13 heading ---
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith("13.") and "실행 로드맵" in text:
        insert_idx = i
        break

if insert_idx is None:
    raise RuntimeError("Could not find Chapter 13 heading to insert before")

print(f"Inserting Chapter 12 before paragraph {insert_idx}: '{doc.paragraphs[insert_idx].text[:50]}'")

# --- Helper: get style references from existing doc ---
# Find the heading style used by other chapters
h1_style = None
h2_style = None
h3_style = None
normal_style = None
for p in doc.paragraphs:
    if p.style and p.style.name:
        if 'Heading 1' in p.style.name and h1_style is None:
            h1_style = p.style.name
        elif 'Heading 2' in p.style.name and h2_style is None:
            h2_style = p.style.name
        elif 'Heading 3' in p.style.name and h3_style is None:
            h3_style = p.style.name

h1_style = h1_style or 'Heading 1'
h2_style = h2_style or 'Heading 2'
h3_style = h3_style or 'Heading 3'

print(f"Using styles: H1={h1_style}, H2={h2_style}, H3={h3_style}")

# --- Build Chapter 12 content ---
chapter_content = []

def add(style, text):
    chapter_content.append((style, text))

# Main heading
add(h1_style, "12.  마케팅 및 콘텐츠 운영 전략")

add('Normal', "본 장은 양면 시장의 '콜드스타트' 문제를 해결하고, 런칭 전·후의 콘텐츠 시딩 및 마케팅 채널 전략을 정의한다. 핵심 원칙: 콘텐츠 먼저 → 트래픽 확보 → 셀러 유치 → 양면 효과 작동.")

# 12.1
add(h2_style, "12.1 콜드스타트 대응 전략")

add('Normal', "양면 시장의 핵심 문제는 '셀러가 없으면 바이어가 안 오고, 바이어가 없으면 셀러가 안 온다'는 것이다. 이를 콘텐츠 시딩으로 해결한다.")

add(h3_style, "A. 유튜브 영상 큐레이션 (비용 0, 최우선)")

add('Normal', "셀러 가입 전에도 YouTube Data API로 카테고리별 영상을 자동 수집하여 Watch 페이지와 카테고리 랜딩을 채운다. 런칭 전 목표: K-뷰티 50편, K-푸드 40편, K-패션 30편, K-컬처 20편, B2B 20편 등 총 160편 이상. 에디터가 품질 검수 후 카테고리·태그를 부여하고, 각 영상에 '관련 상품 찾기' CTA를 연결한다.")

add(h3_style, "B. 브랜드 프로필 사전 구축 (프로필 클레임 모델)")

add('Normal', "이미 미국에 진출한 한국 브랜드의 프로필을 공개 정보 기반으로 운영팀이 직접 생성한다. Tier 1: 대형 K-뷰티 브랜드(조선미녀, 코스알엑스 등) 30개, Tier 2: 아마존·큐텐 입점 중소 브랜드 50개, Tier 3: 수출바우처 참여 기업 50개. 총 130개 프로필에 '이 브랜드의 공식 담당자이신가요? → 프로필 클레임하기' CTA를 배치하여 셀러 가입을 유도한다(LinkedIn Company Page 모델). 예상 클레임 전환율 10~15%.")

add(h3_style, "C. 매거진/에디토리얼 콘텐츠")

add('Normal', "SEO 유입과 사이트 전문성 확보를 위한 읽을거리. K-뷰티 가이드(주 2회), 바이어 가이드(주 1회, 예: 'How to Import Korean Cosmetics to the US: FDA Guide'), 트렌드 리포트(격주), 셀러 성공사례(월 2회), 영상 위클리(주 1회). 런칭 전 20편, 런칭 후 주 3~4편 발행.")

add(h3_style, "D. 카테고리 랜딩 페이지")

add('Normal', "8개 대카테고리 각각에 카테고리 설명, 시장 규모 요약(기획서 조사 데이터 활용), 큐레이션 영상 6~8편, 시딩된 브랜드 프로필 카드, '이 카테고리에 입점하기' 셀러 CTA를 사전 배치한다. 빈 카테고리 0개를 목표로 한다.")

# 12.2
add(h2_style, "12.2 셀러 확보 마케팅")

add('Normal', "셀러 유치 5대 채널:")

add('Normal', "▪ 채널 1: 프로필 클레임 전환 — 시딩된 130개 브랜드에 클레임 CTA. 클레임 시 Pro 기능 1개월 무료 제공.")

add('Normal', "▪ 채널 2: 정부 수출지원사업 연계 — 수출바우처(4,000개사, 1,502억원) 수행기관 등록, 중기부 해외마케팅·KOTRA 수출상담회·지자체 수출지원 프로그램 제휴. 바우처 예산으로 유료 입점이 가능해져 셀러 비용 부담 해소.")

add('Normal', "▪ 채널 3: 한인 비즈니스 커뮤니티 직접 영업 — LA·NY 한인 상공회의소(각 300~500개사), 한인 무역협회(1,000+ 기업), K-뷰티 셀러 카페·단톡(5,000+ 셀러), 아마존 한국 셀러 그룹(10,000+ 셀러) 대상 무료 입점 이벤트.")

add('Normal', "▪ 채널 4: 콘텐츠 마케팅(인바운드) — FDA 가이드, 시장 트렌드 리포트, 수출바우처 활용법 등 셀러가 검색으로 찾아오는 콘텐츠. CTA로 입점 연결.")

add('Normal', "▪ 채널 5: 온보딩 인센티브 — 얼리버드 무료 Pro 6개월(첫 100개사), 프리미엄 셀러 영상 제작 지원, 레퍼럴 크레딧(셀러→셀러 추천 시 1개월 무료), 월 2건 성공사례 피처링.")

add('Normal', "셀러 확보 목표: 런칭 시점 50개사 → 3개월 200개사(유료 20) → 6개월 500개사(유료 75) → 12개월 1,000개사(유료 200).")

# 12.3
add(h2_style, "12.3 바이어/소비자 유입 마케팅")

add('Normal', "바이어 유입 5대 채널:")

add('Normal', "▪ 채널 1: SEO/콘텐츠 마케팅(장기 핵심) — 카테고리 SEO(롱테일 키워드 최적화), 매거진 SEO(주 3~4편, 구매·수입 가이드), 상품/업체 페이지 구조화 데이터(Product·VideoObject), 영상 SEO(영상 사이트맵). 타깃 키워드: 'buy korean cosmetics wholesale', 'best korean skincare 2026' 등.")

add('Normal', "▪ 채널 2: YouTube 자체 채널 — 상품 리뷰/비교(주 2회), 'Top 10 Korean [카테고리]'(월 2회), 셀러 인터뷰/공장 투어(월 2회), Shorts(주 3회). 목표: 6개월 구독 5,000 / 12개월 10,000.")

add('Normal', "▪ 채널 3: SNS 마케팅 — Instagram(K-웨이브 소비자 18-35, 일 1회), TikTok(Gen Z, 일 1회), LinkedIn(B2B 바이어, 주 3회), Facebook(재외 한인 35+, 주 5회), X(K-컬처 팬덤, 일 2회).")

add('Normal', "▪ 채널 4: 뉴스레터 — K-Trend Weekly(바이어/소비자용 주 1회), Seller Digest(셀러용 격주), Buyer Brief(B2B 바이어용 월 2회). 목표: 6개월 내 구독자 5,000명.")

add('Normal', "▪ 채널 5: 유료 광고(Phase 2부터) — 런칭 초기 3개월은 유기 채널에 집중, 이후 Google Ads 검색(40%), Meta Ads(25%), YouTube Ads(20%), LinkedIn Ads(15%).")

add('Normal', "바이어 유입 목표: 1개월 5,000 UV(해외 40%) → 3개월 20,000 UV(50%) → 6개월 50,000 UV(55%) → 12개월 150,000 UV(60%+).")

# 12.4
add(h2_style, "12.4 콘텐츠 운영 체계")

add('Normal', "콘텐츠 운영은 세 축으로 구성한다:")

add('Normal', "▪ 자동화 콘텐츠: YouTube Data API를 통한 영상 자동 수집·정렬(조회수·최신순), 신규 셀러·상품 등록 시 자동 생성되는 프로필 페이지, 트렌드 데이터 기반 자동 큐레이션.")

add('Normal', "▪ 에디토리얼 콘텐츠: 에디터가 직접 작성하는 매거진·가이드·리포트. AI 보조 작성 + 에디터 검수로 생산성 3배 향상. 주 3~4편 발행.")

add('Normal', "▪ 셀러 생성 콘텐츠(UGC): 셀러가 직접 등록하는 상품 정보, 영상 URL, 브랜드 스토리. 운영팀 검수 프로세스를 거쳐 공개.")

add('Normal', "콘텐츠 품질 관리: 영상은 조회수·좋아요 비율·댓글 품질로 자동 필터링 후 에디터 최종 승인. 매거진은 출처 명시 의무화, 팩트체크 프로세스 적용. 셀러 콘텐츠는 사업자등록 확인·인증 배지 기준에 따라 등급 부여.")

# 12.5
add(h2_style, "12.5 시딩 타임라인 및 런칭 시점 목표")

add('Normal', "런칭 8주 전부터 시딩 시작: W1-2 YouTube 큐레이션 160편, W2-3 Tier 1 프로필 30개, W3-4 카테고리 랜딩 8개, W4-5 Tier 2-3 프로필 100개, W5-6 매거진 20편, W6-7 QA + SEO 검수, W7-8 소프트 런칭(한인 커뮤니티 대상).")

add('Normal', "런칭 시점 사이트 상태 목표: 큐레이션 영상 160편+, 브랜드 프로필 130개+, 매거진 콘텐츠 20편+, 카테고리 완성 8/8, 인덱싱 대상 페이지 500+. 빈 페이지 0개, 첫 화면에서 5초 내 가치 전달.")

add('Normal', "월간 운영 비용(Phase 1 기준): 콘텐츠 에디터 1명, SNS/마케팅 1명, 영상 외주, 인프라 비용 포함 약 $7,200~10,500/월. Phase 2(유료 광고 포함) 약 $9,200~15,500/월. 수출바우처 수행기관 등록 시 마케팅 비용 일부 충당 가능.")


# --- Insert content into document ---
# We insert from bottom to top so indices stay valid
from docx.oxml.ns import qn
from lxml import etree

body = doc.element.body
# Get the XML element of paragraph at insert_idx
ref_element = doc.paragraphs[insert_idx]._element

# Insert in forward order; addprevious always goes immediately before ref
for style, text in chapter_content:
    new_p = deepcopy(doc.paragraphs[0]._element)  # clone a paragraph
    # Clear cloned content
    for child in list(new_p):
        new_p.remove(child)
    
    # Create fresh paragraph using python-docx API by adding to end, then moving
    temp_p = doc.add_paragraph(text, style=style)
    new_el = temp_p._element
    
    # Move before the reference element
    body.remove(new_el)
    ref_element.addprevious(new_el)

# --- Also update TOC entry: add "12. 마케팅 및 콘텐츠 운영 전략" to table of contents ---
# Find the TOC area (between "목차" and "0. 기획 개요")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "11." in text and "기술 스택" in text and "다국어" in text:
        # Check if this is a TOC entry (near the beginning)
        if i < 55:  # TOC is in the first ~50 paragraphs
            # Find the next paragraph (should be "13. 실행 로드맵" in TOC)
            toc_insert_ref = doc.paragraphs[i + 1]._element
            
            # Create TOC entry paragraph
            toc_p = doc.add_paragraph("12.  마케팅 및 콘텐츠 운영 전략\t19", style=p.style.name if p.style else 'Normal')
            toc_el = toc_p._element
            body.remove(toc_el)
            toc_insert_ref.addprevious(toc_el)
            print(f"Added TOC entry after paragraph {i}")
            break

# Save
doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")
print(f"Chapter 12 inserted with {len(chapter_content)} paragraphs")
