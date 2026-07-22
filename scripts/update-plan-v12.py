#!/usr/bin/env python3
"""
LoveKorea.Us 기획서 v1.1 → v1.2 업데이트
추가 내용:
  - 12장 뒤에: 유저 플로우 다이어그램 (이미지 2장)
  - 그 뒤에: 디자인 시안 (이미지 8장)
  - 마케팅/콘텐츠 전략 (기존 12장에 이미 포함되어 있으므로, 보완 내용 추가)
  
접근: python-docx로 기존 문서 끝(부록 전)에 새 챕터 삽입
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil

# Paths
src = os.path.expanduser('~/Downloads/LoveKoreaUs_홈페이지_기획서_v1.1.docx')
dst = os.path.expanduser('~/Downloads/LoveKoreaUs_홈페이지_기획서_v1.2.docx')

# Copy first
shutil.copy2(src, dst)

doc = Document(dst)

NAVY = RGBColor(0x0B, 0x2E, 0x59)
RED = RGBColor(0xC8, 0x20, 0x2F)
GRAY = RGBColor(0x66, 0x66, 0x66)

# ── Find insertion point: before "부록 · 조사 출처" ──
insert_before_idx = None
for i, p in enumerate(doc.paragraphs):
    if '부록' in p.text and '조사 출처' in p.text:
        insert_before_idx = i
        break

if insert_before_idx is None:
    # fallback: insert before last paragraph
    insert_before_idx = len(doc.paragraphs) - 1

print(f"Insert point: paragraph {insert_before_idx} ('{doc.paragraphs[insert_before_idx].text[:40]}')")

# ── Helper: insert content before a specific paragraph ──
# python-docx doesn't have insert_before natively, so we'll append to the end
# and the content will appear after existing chapters but we'll use addprevious
from lxml import etree

ref_element = doc.paragraphs[insert_before_idx]._element

def add_heading_before(text, level=1):
    """Add a heading paragraph before the reference element."""
    p = doc.add_paragraph()  # temporary at end
    # Style it
    p.style = doc.styles[f'Heading {level}']
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    # Move before reference
    ref_element.addprevious(p._element)
    return p

def add_para_before(text, bold=False, italic=False, size=Pt(10), color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    ref_element.addprevious(p._element)
    return p

def add_image_before(image_path, width_inches=6.0, caption=None):
    """Add an image paragraph before the reference."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    ref_element.addprevious(p._element)
    
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.size = Pt(9)
        cr.font.color.rgb = GRAY
        cr.italic = True
        ref_element.addprevious(cp._element)
    
    # Add spacing
    sp = doc.add_paragraph()
    ref_element.addprevious(sp._element)
    return p

def add_pagebreak_before():
    p = doc.add_paragraph()
    run = p.add_run()
    from docx.oxml.ns import qn
    br = etree.SubElement(run._element, qn('w:br'))
    br.set(qn('w:type'), 'page')
    ref_element.addprevious(p._element)

# ── Image paths ──
design_images = [
    ('/Users/martin/.verdent/generate_images/generate_20260722_230832_76a921c8.png', '시안 1: 홈페이지 (Homepage) — Desktop'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_230832_8e424865_1.png', '시안 2: For Sellers 랜딩 페이지'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_230832_ce34a875_2.png', '시안 3: Watch (영상 큐레이션) 페이지'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_230832_7ce8b65e_3.png', '시안 4: Directory (업소록) 페이지'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_231144_1d59924f.png', '시안 5: 상품 상세 (Product Detail) 페이지'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_231144_e0b480e6_1.png', '시안 6: 업체/브랜드 상세 페이지'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_231144_94b620de_2.png', '시안 7: Deals & Promotions 페이지'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_231144_bf6236e3_3.png', '시안 8: 모바일 홈 (Mobile Home)'),
]

flow_images = [
    ('/Users/martin/.verdent/generate_images/generate_20260722_231315_5313a7a7.png', '셀러 온보딩 플로우 (Seller Onboarding Flow)'),
    ('/Users/martin/.verdent/generate_images/generate_20260722_231315_67864fd4_1.png', '바이어 탐색 플로우 (Buyer Discovery Flow)'),
]


# ═══════════════════════════════════════════════
# CHAPTER: 유저 플로우 다이어그램
# ═══════════════════════════════════════════════
add_pagebreak_before()

add_heading_before('16.  유저 플로우 다이어그램 (User Flow Diagrams)', level=1)
add_para_before('본 플랫폼의 핵심 사용자 여정을 다이어그램으로 정의한다. 셀러(공급측)와 바이어(수요측) 두 축의 플로우를 각각 시각화하여, 개발 시 UX 설계의 기준으로 활용한다.')

add_heading_before('16.1 셀러 온보딩 플로우', level=2)
add_para_before('한국 기업/소상공인이 플랫폼에 가입하고, 프로필을 작성하며, 상품을 등록하여 해외 바이어로부터 문의를 받기까지의 전체 여정이다.')
add_para_before('핵심 단계: 회원가입 → 셀러 등록 신청 → 프로필/상품 입력 → 심사(관리자) → 승인 → 문의 수신 → 유료 전환', bold=True, size=Pt(9), color=GRAY)
add_image_before(flow_images[0][0], 5.0, flow_images[0][1])

add_heading_before('16.2 바이어 탐색 플로우', level=2)
add_para_before('해외 바이어/소비자가 검색·SNS·유튜브 등 다양한 경로로 유입되어, 상품/업체를 발견하고 신뢰를 확인한 뒤 문의·구매로 전환되는 여정이다.')
add_para_before('핵심 단계: 유입(검색/SNS/YouTube) → 탐색(카테고리/디렉토리/Watch) → 상세 확인(영상·리뷰·인증) → 행동(문의/견적/구매) → 리뷰', bold=True, size=Pt(9), color=GRAY)
add_image_before(flow_images[1][0], 5.0, flow_images[1][1])


# ═══════════════════════════════════════════════
# CHAPTER: 마케팅/콘텐츠 운영 전략 (보완)
# ═══════════════════════════════════════════════
add_pagebreak_before()

add_heading_before('17.  초기 콘텐츠 시딩 & 마케팅 전략 (Content Seeding & Marketing)', level=1)
add_para_before('양면 시장의 콜드스타트 문제를 해결하고, 런칭 시점부터 사이트가 "비어 보이지 않도록" 하기 위한 콘텐츠 시딩 전략과 단계별 마케팅 계획을 정의한다.')

add_heading_before('17.1 콜드스타트 해법', level=2)
add_para_before('해법 순서: 콘텐츠 먼저 → 트래픽 확보 → 셀러 유치 → 양면 효과 작동', bold=True)

add_heading_before('17.2 시딩 콘텐츠 유형 및 목표', level=2)
seeding = [
    'A. 유튜브 영상 큐레이션 (비용 0) — 런칭 전 160편 목표. YouTube Data API로 키워드 기반 자동 수집 + 에디터 품질 검수.',
    'B. 업체/브랜드 프로필 사전 구축 — 공개 정보 기반 130개 프로필 생성. "공식 담당자이신가요? → 프로필 클레임" CTA로 셀러 전환 (예상 10~15%).',
    'C. 매거진/에디토리얼 — SEO 유입 + 전문성 확보. 런칭 전 20편, 이후 주 3~4편. "Best Korean Sunscreens 2026", "How to Import Korean Cosmetics" 등.',
    'D. 카테고리 랜딩 — 8개 대분류 각각에 설명 + 트렌드 요약 + 큐레이션 영상 6~8편 + 셀러 CTA.',
]
for s in seeding:
    add_para_before(s, size=Pt(9))

add_heading_before('17.3 셀러 확보 채널 (5대 채널)', level=2)
seller_channels = [
    '① 프로필 클레임 전환 — 시딩된 130개 프로필의 "공식 담당자" CTA → 셀러 가입. 클레임 시 Pro 1개월 무료.',
    '② 정부 수출지원사업 연계 — 수출바우처(4,000개사/1,502억원), 중기부 해외마케팅, KOTRA 수출상담회 등.',
    '③ 한인 비즈니스 커뮤니티 — LA·NY 한인 상공회의소, 한인 무역협회, K-뷰티 셀러 커뮤니티.',
    '④ 콘텐츠 마케팅 (인바운드) — FDA 가이드, 시장 리포트 등 매거진으로 셀러가 직접 유입.',
    '⑤ 온보딩 인센티브 — 얼리버드 무료 Pro(6개월), 영상 제작 지원, 레퍼럴 크레딧.',
]
for ch in seller_channels:
    add_para_before(ch, size=Pt(9))

add_heading_before('17.4 바이어 유입 채널 (5대 채널)', level=2)
buyer_channels = [
    '① SEO/콘텐츠 마케팅 — 카테고리 SEO, 매거진, 구조화 데이터. 6개월 내 월 5만 UV 유기 유입 목표.',
    '② YouTube 자체 채널 — 상품 리뷰 주 2회, Top 10 시리즈 월 2회, Shorts 주 3회. 12개월 내 구독 1만.',
    '③ SNS — Instagram(일 1), TikTok(일 1), LinkedIn(주 3), Facebook(주 5), X(일 2).',
    '④ 뉴스레터 — K-Trend Weekly(주 1), Seller Digest(격주), Buyer Brief(월 2). 6개월 내 5,000 구독.',
    '⑤ 유료 광고 (Phase 2 이후) — Google Ads(40%), Meta(25%), YouTube(20%), LinkedIn(15%).',
]
for ch in buyer_channels:
    add_para_before(ch, size=Pt(9))

add_heading_before('17.5 KPI 마일스톤', level=2)
kpi_text = """런칭 시점: 영상 160+ / 프로필 130+ / 매거진 20+
+3개월: 셀러 200 / 월 UV 2만 / 리드 200건 / MRR $2K
+6개월: 셀러 500 / 월 UV 5만 / 리드 800건 / MRR $10K
+12개월: 셀러 1,000 / 유료 200 / 월 UV 15만 / 리드 3,000건 / MRR $40K"""
add_para_before(kpi_text, size=Pt(9))

add_heading_before('17.6 예산 프레임', level=2)
budget_text = """Phase 1 (유기 중심): $7,200~10,500/월 — 에디터 1명 + 마케팅 1명 + 영상 외주 + 인프라
Phase 2 (유료 확대): $9,200~15,500/월 — 유료 광고 $2,000~5,000 추가
비용 절감: 수출바우처 활용, AI 보조 콘텐츠, YouTube API 자동화, 인플루언서 교환 제휴"""
add_para_before(budget_text, size=Pt(9))


# ═══════════════════════════════════════════════
# CHAPTER: 디자인 시안
# ═══════════════════════════════════════════════
add_pagebreak_before()

add_heading_before('18.  디자인 시안 (Design Mockups)', level=1)
add_para_before('기획서 6장(페이지별 구성)과 10장(디자인/UX 가이드)을 기반으로 제작한 8장의 디자인 시안이다. 컬러 시스템(네이비 #0B2E59 + 레드 #C8202F), 카드형 레이아웃, 영상 중심 비주얼 언어를 적용했다.')

for img_path, caption in design_images:
    # Check if mobile (last image)
    is_mobile = '모바일' in caption or 'Mobile' in caption
    width = 4.0 if is_mobile else 6.2
    
    add_heading_before(caption, level=2)
    add_image_before(img_path, width, caption)


# ── Save ──
doc.save(dst)

# Verify
doc2 = Document(dst)
print(f"\n=== 저장 완료 ===")
print(f"파일: {dst}")
print(f"총 단락: {len(doc2.paragraphs)}")
print(f"총 테이블: {len(doc2.tables)}")

# Count new chapters
new_chapters = [p.text for p in doc2.paragraphs if p.text.startswith('16.') or p.text.startswith('17.') or p.text.startswith('18.')]
print(f"\n새 챕터:")
for ch in new_chapters:
    print(f"  {ch[:60]}")
